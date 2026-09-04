"""
Module: html_generator.py
Created: 2026-09-03
Purpose: Renders template-ready resume data into HTML and safe PDF/DOCX
         formats using Jinja2, Playwright (Chromium) for PDF, and
         python-docx for DOCX.
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from jinja2.sandbox import SandboxedEnvironment

from app.config import settings
from app.utils.exceptions import TemplateRenderError


def render_html(
    template_source: str,
    data: dict,
    *,
    sandbox: bool = True,
    loader_dir: Optional[str] = None,
) -> str:
    """Render resume data to an HTML string.

    Args:
        template_source: Jinja2 HTML template source.
        data: Template-ready resume data.
        sandbox: Whether to use SandboxedEnvironment (required for custom).
        loader_dir: Optional directory root for resolving {% extends %}.

    Returns:
        str: The fully rendered HTML.

    Raises:
        TemplateRenderError: If rendering fails or the template is unsafe.
    """
    try:
        from jinja2 import Environment
        from jinja2.loaders import FileSystemLoader

        if sandbox:
            loader = FileSystemLoader(loader_dir) if loader_dir else None
            env = SandboxedEnvironment(
                autoescape=True, loader=loader, enable_async=False
            )
        else:
            env = Environment(
                autoescape=True,
                loader=FileSystemLoader(loader_dir) if loader_dir else None,
            )
        template = env.from_string(template_source)
        return template.render(**data)
    except Exception as exc:
        raise TemplateRenderError(f"HTML render failed: {exc}") from exc


def render_pdf(
    template_source: str,
    data: dict,
    config: dict,
    output_path: Path,
    *,
    sandbox: bool = True,
    loader_dir: Optional[str] = None,
) -> Path:
    """Render resume data to a PDF file.

    Args:
        template_source: Jinja2 HTML template source.
        data: Template-ready resume data.
        config: Template layout configuration.
        output_path: Destination for the generated PDF.
        sandbox: Whether to sandbox the Jinja environment.
        loader_dir: Optional loader root for {% extends %}.

    Returns:
        Path: The path to the generated PDF.
    """
    html = render_html(template_source, data, sandbox=sandbox, loader_dir=loader_dir)
    _html_to_pdf(html, config, output_path)
    return output_path


def _html_to_pdf(html: str, config: dict, output_path: Path) -> None:
    """Render the exact preview HTML to a styled PDF via Chromium.

    Uses the same rendered HTML+CSS string the browser preview shows, so the
    downloaded PDF reproduces the template's true layout (two-column sidebars,
    flex alignment, colored bands, skill chips) instead of a flattened generic
    document.

    Args:
        html: The fully rendered template HTML (identical to the preview).
        config: Template layout config (kept for API compatibility).
        output_path: Destination PDF path.

    Raises:
        TemplateRenderError: If Chromium cannot be launched or the PDF fails.
    """
    try:
        import asyncio
        import sys

        # On Windows, Playwright's sync API builds its own event loop and calls
        # asyncio.create_subprocess_exec to launch Chromium. A Selector loop
        # raises NotImplementedError there, so force the Proactor policy in this
        # thread before Playwright starts. This is thread-local-scoped and avoids
        # disturbing the app's serving loop or the uvicorn --reload reloader.
        if sys.platform == "win32":
            asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

        from playwright.sync_api import sync_playwright

        # Playwright's sync API is bound to the calling thread. Launch a fresh
        # browser per call so concurrent worker threads and direct calls each
        # own their own instance (avoids "Cannot switch to a different thread").
        with sync_playwright() as p:
            browser = p.chromium.launch()
            try:
                page = browser.new_page()
                page.set_content(html, wait_until="load")
                page.pdf(
                    path=str(output_path),
                    format="Letter",
                    print_background=True,
                    prefer_css_page_size=False,
                )
            finally:
                browser.close()
    except Exception as exc:
        detail = f"{type(exc).__name__}: {exc!r}"
        raise TemplateRenderError(f"PDF render failed: {detail}") from exc


def _extract_headings(html: str) -> list[str]:
    """Find heading texts within h1-h3 tags.

    Args:
        html: Rendered HTML.

    Returns:
        list[str]: Heading texts found.
    """
    return re.findall(r"<(?:h[1-3])[^>]*>(.*?)</(?:h[1-3])>", html, re.S | re.I)


def _strip_tags(html: str) -> str:
    """Remove HTML tags, turning block boundaries into newlines.

    Drops the inner content of ``<style>``/``<script>`` blocks and the whole
    ``<head>`` region first, so CSS rules, scripts, and metadata never leak
    into the extracted text used for PDF/DOCX generation.

    Args:
        html: The HTML content.

    Returns:
        str: Text with tags stripped and paragraphs separated by newlines.
    """
    html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.I | re.S)
    html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.I | re.S)
    html = re.sub(r"<head[^>]*>.*?</head>", "", html, flags=re.I | re.S)
    html = re.sub(r"<(?:li|p|div|br|tr)[^>]*>", "\n", html, flags=re.I)
    html = re.sub(r"</(?:li|p|div|tr)[^>]*>", "\n", html, flags=re.I)
    html = re.sub(r"<[^>]+>", "", html)
    html = re.sub(r"\n{2,}", "\n", html)
    import html as htmlmod

    return htmlmod.unescape(html).strip()


def _split_sections(text: str, headings: list[str]) -> list[tuple[Optional[str], list[str]]]:
    """Split body text into (heading, lines) tuples by known headings.

    Args:
        text: Tag-stripped body text.
        headings: Known heading texts.

    Returns:
        list of (heading or None, lines) sections.
    """
    lines = [ln.strip() for ln in text.splitlines()]
    sections: list[tuple[Optional[str], list[str]]] = []
    current_title: Optional[str] = None
    current_lines: list[str] = []

    for line in lines:
        if not line:
            continue
        if line in headings:
            if current_lines or current_title:
                sections.append((current_title, current_lines))
            current_title = line
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines or current_title:
        sections.append((current_title, current_lines))
    return sections





def render_docx(
    template_source: str,
    data: dict,
    config: dict,
    output_path: Path,
    *,
    sandbox: bool = True,
    loader_dir: Optional[str] = None,
) -> Path:
    """Render resume data to a DOCX file, preserving content.

    Args:
        template_source: Jinja2 HTML template source (used for structure).
        data: Template-ready resume data.
        config: Template layout configuration.
        output_path: Destination for the generated DOCX.
        sandbox: Whether to sandbox the Jinja environment.
        loader_dir: Optional loader root for {% extends %}.

    Returns:
        Path: The path to the generated DOCX.
    """
    html = render_html(template_source, data, sandbox=sandbox, loader_dir=loader_dir)
    _html_to_docx(html, config, output_path)
    return output_path


def _html_to_docx(html: str, config: dict, output_path: Path) -> None:
    """Build a DOCX from the rendered HTML structure.

    Args:
        html: Rendered HTML with resume content.
        config: Layout config providing fonts/colors.
        output_path: Destination DOCX path.
    """
    doc = DocxDocument()
    color_hex = config.get("accent_color", "#2563EB")
    accent = _hex_to_rgb(color_hex)

    headings = _extract_headings(html)
    body = _strip_tags(html)
    sections = _split_sections(body, headings)

    for title, lines in sections:
        if title:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = accent
        first = True
        for line in lines:
            if not line.strip():
                continue
            para = doc.add_paragraph()
            if not first and line.strip().startswith(("-", "•")):
                para.style = doc.styles["List Bullet"]
            run = para.add_run(_strip_list_marker(line))
            run.font.size = Pt(10)
            first = False

    doc.save(str(output_path))


def _strip_list_marker(line: str) -> str:
    """Remove leading bullet markers from a line.

    Args:
        line: Text line.

    Returns:
        str: Line without a leading '-' or '•' marker.
    """
    return line.lstrip("-• ")


def _hex_to_rgb(color_hex: str) -> RGBColor:
    """Convert a hex color string to a docx RGBColor.

    Args:
        color_hex: Hex string like '#2563EB'.

    Returns:
        RGBColor: The parsed color.
    """
    color_hex = color_hex.lstrip("#")
    if len(color_hex) != 6:
        color_hex = "2563EB"
    r = int(color_hex[0:2], 16)
    g = int(color_hex[2:4], 16)
    b = int(color_hex[4:6], 16)
    return RGBColor(r, g, b)
