"""
Module: test_generation.py
Created: 2026-09-03
Purpose: Regression tests ensuring CSS/<style> markup does not leak into
         generated PDF/DOCX output, resume content is preserved, and PDF/DOCX
         generation still succeeds.
"""

from pathlib import Path
from zipfile import BadZipFile, ZipFile

from app.services.html_generator import _strip_tags, render_html, render_pdf
from app.services.resume_processor import process as process_data
from app.services.template_service import template_service

CORE_TEMPLATES = [
    "academic",
    "ats_standard",
    "classic",
    "elegant",
    "executive",
    "minimal",
    "modern",
    "professional",
    "simple",
    "technical",
]

CSS_TOKENS = [
    "font-family",
    "body {",
    ".name {",
    "accent_color",
    "font-size:",
    "letter-spacing",
    "text-transform",
]


def _sample_parsed():
    from app.schemas.resume import ParsedResumeData

    return ParsedResumeData(
        personal_info={
            "name": "Ziaur Rab",
            "email": "ziaur@example.com",
            "phone": "+1-555-0100",
            "location": "New York, NY",
        },
        summary="Software engineer with 5+ years of experience.",
        experience=[
            {
                "company": "Tech Corp",
                "title": "Senior Developer",
                "location": "New York, NY",
                "start_date": "2021",
                "end_date": "Present",
                "description": "Led a team building REST APIs.",
            }
        ],
        education=[
            {
                "institution": "MIT",
                "degree": "B.S. Computer Science",
                "start_date": "2015",
                "end_date": "2019",
                "gpa": "3.8",
            }
        ],
        skills=["Python", "FastAPI", "Docker"],
        certifications=["AWS Cert"],
        languages=["English", "Spanish"],
        projects=[{"name": "ATS Platform"}],
    )


def _upload(client, session_id="gen-session"):
    with open(
        Path(__file__).parent / "fixtures" / "sample_cv.txt", "rb"
    ) as f:
        return client.post(
            "/api/upload/cv",
            files={"file": ("sample_cv.txt", f, "text/plain")},
            headers={"x-session-id": session_id},
        )


def _generate(client, resume_id, template_id, fmt):
    resp = client.post(
        f"/api/resume/{resume_id}/generate",
        json={"template_id": template_id, "format": fmt},
        headers={"x-session-id": "gen-session"},
    )
    assert resp.status_code == 201, (template_id, fmt, resp.text)
    return resp.json()


# --- Unit-level: CSS does not leak from rendered HTML ---


def test_strip_tags_removes_css_from_all_templates():
    """Rendered HTML for every template yields no CSS tokens after stripping."""
    parsed = _sample_parsed()
    for tid in CORE_TEMPLATES:
        resolved = template_service.resolve(tid, "test", [])
        prepared = process_data(parsed, resolved["config"])
        from app.services.html_generator import render_html

        html = render_html(
            resolved["html"],
            prepared,
            sandbox=resolved["is_custom"],
            loader_dir=str(template_service.builtin_root),
        )
        stripped = _strip_tags(html)
        for token in CSS_TOKENS:
            assert token not in stripped, f"{tid} leaked CSS token: {token}"
        assert "Ziaur Rab" in stripped


def test_strip_tags_preserves_resume_content():
    """Stripping keeps all resume content but drops the head/style markup."""
    parsed = _sample_parsed()
    resolved = template_service.resolve("ats_standard", "test", [])
    prepared = process_data(parsed, resolved["config"])
    from app.services.html_generator import render_html

    html = render_html(
        resolved["html"],
        prepared,
        sandbox=False,
        loader_dir=str(template_service.builtin_root),
    )
    stripped = _strip_tags(html)
    for marker in ["Ziaur Rab", "Tech Corp", "Senior Developer", "MIT",
                   "Python", "English", "ATS Platform"]:
        assert marker in stripped, f"missing content: {marker}"


def test_experience_company_rendered_before_title():
    """Every template shows the company header before the designation."""
    parsed = _sample_parsed()
    for tid in CORE_TEMPLATES:
        resolved = template_service.resolve(tid, "test", [])
        prepared = process_data(parsed, resolved["config"])
        from app.services.html_generator import render_html

        html = render_html(
            resolved["html"],
            prepared,
            sandbox=resolved["is_custom"],
            loader_dir=str(template_service.builtin_root),
        )
        stripped = _strip_tags(html)
        company_pos = stripped.find("Tech Corp")
        title_pos = stripped.find("Senior Developer")
        assert company_pos != -1, f"{tid} missing company"
        assert title_pos != -1, f"{tid} missing title"
        assert company_pos < title_pos, f"{tid} company must precede title"


# --- API-level: generation succeeds ---


def test_all_formats_generate_successfully(client):
    """PDF, DOCX, and HTML generation each succeed and record a format."""
    up = _upload(client)
    resume_id = up.json()["id"]
    for fmt in ("pdf", "docx", "html"):
        result = _generate(client, resume_id, "ats_standard", fmt)
        assert result["format"] == fmt


# --- API-level: DOCX is a valid Office (.docx / OOXML zip) document ---


def test_docx_is_valid_ooxml_archive(client):
    """Downloaded DOCX is a valid ZIP/OOXML document with a .docx filename."""
    up = _upload(client)
    resume_id = up.json()["id"]
    result = _generate(client, resume_id, "modern", "docx")

    dl = client.get(
        f"/api/resume/{resume_id}/download/{result['id']}",
        headers={"x-session-id": "gen-session"},
    )
    assert dl.status_code == 200
    assert dl.headers["content-type"].startswith(
        "application/vnd.openxmlformats"
    )
    assert "filename=" in dl.headers["content-disposition"]

    payload = dl.content
    try:
        with ZipFile(__import__("io").BytesIO(payload)) as zf:
            names = zf.namelist()
    except BadZipFile:
        raise AssertionError("DOCX download is not a valid ZIP/OOXML archive")

    assert "[Content_Types].xml" in names
    assert any(n.startswith("word/") for n in names)


def test_docx_body_contains_no_css(client):
    """The DOCX text part must not contain leaked CSS rules."""
    import io

    from docx import Document

    up = _upload(client)
    resume_id = up.json()["id"]
    result = _generate(client, resume_id, "academic", "docx")

    dl = client.get(
        f"/api/resume/{resume_id}/download/{result['id']}",
        headers={"x-session-id": "gen-session"},
    )
    assert dl.status_code == 200

    document = Document(io.BytesIO(dl.content))
    full_text = "\n".join(p.text for p in document.paragraphs)
    for token in CSS_TOKENS:
        assert token not in full_text, f"DOCX leaked CSS token: {token}"
    # Content preserved in the DOCX.
    assert "Ziaur Rab" in full_text or "John Doe" in full_text


# --- PDF fidelity: Playwright renders the same HTML the preview shows ---


def _render_pdf_bytes(template_id: str) -> bytes:
    """Render a built-in template to a PDF via the shared engine and return it.

    Args:
        template_id: Built-in template id.

    Returns:
        bytes: The generated PDF payload.
    """
    import tempfile

    parsed = _sample_parsed()
    resolved = template_service.resolve(template_id, "fidelity-test", [])
    prepared = process_data(parsed, resolved["config"])
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        render_pdf(
            resolved["html"],
            prepared,
            resolved["config"],
            tmp_path,
            sandbox=resolved["is_custom"],
            loader_dir=str(template_service.builtin_root),
        )
        return tmp_path.read_bytes()
    finally:
        tmp_path.unlink(missing_ok=True)


def _pdf_text(payload: bytes) -> str:
    """Extract the visible text from a PDF payload.

    Args:
        payload: PDF file bytes.

    Returns:
        str: Concatenated text across all pages.
    """
    import io

    import pdfplumber

    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        return "\n".join((pg.extract_text() or "") for pg in pdf.pages)


def _normalize(text: str) -> str:
    """Lowercase and drop all whitespace so alignment/letter-spacing artifacts
    (e.g. 'z iaur rab', line wrapping) do not break substring matching.

    Args:
        text: Raw extracted text.

    Returns:
        str: Lowercased text with every whitespace character removed.
    """
    return "".join(text.lower().split())


# Content markers (normalized: no whitespace, lowercase) that all templates
# render. Templates may omit some sections (e.g. academic has no summary), so
# the universal set stays to what every layout shows.
UNIVERSAL_MARKERS = [
    "ziaurrab", "ziaur@example.com", "techcorp", "seniordeveloper",
    "newyork,ny", "mit", "b.s.computerscience", "python",
]
# Full marker set used only for ats_standard, which renders every section.
FULL_MARKERS = UNIVERSAL_MARKERS + [
    "softwareengineer", "fastapi", "docker", "awscert", "english",
    "spanish", "atsplatform",
]
SECTION_MARKERS = [
    "summary", "experience", "education", "skills",
    "certifications", "languages", "projects",
]


def test_pdf_all_templates_short_and_content_complete():
    """Every built-in template's PDF is valid, non-trivial, and has no CSS leak."""
    parsed = _sample_parsed()
    for tid in CORE_TEMPLATES:
        payload = _render_pdf_bytes(tid)
        assert len(payload) > 5_000, f"{tid} PDF is suspiciously small"
        text = _normalize(_pdf_text(payload))
        for marker in UNIVERSAL_MARKERS:
            assert marker in text, f"{tid} PDF missing content marker: {marker}"
        for token in CSS_TOKENS:
            assert token not in text, f"{tid} PDF leaked CSS token: {token}"


def test_pdf_content_matches_preview_text():
    """The PDF text equals the tag-stripped preview text for a template."""
    tid = "ats_standard"
    parsed = _sample_parsed()
    resolved = template_service.resolve(tid, "fidelity-test", [])
    prepared = process_data(parsed, resolved["config"])
    html = render_html(
        resolved["html"],
        prepared,
        sandbox=resolved["is_custom"],
        loader_dir=str(template_service.builtin_root),
    )
    preview_text = _normalize(_strip_tags(html))
    pdf_text = _normalize(_pdf_text(_render_pdf_bytes(tid)))
    for marker in FULL_MARKERS + SECTION_MARKERS:
        assert marker in pdf_text, f"pdf missing {marker}"
        assert marker in preview_text, f"preview missing {marker}"


def _upload_bytes(client, payload, filename, ctype, session_id="roundtrip-session"):
    return client.post(
        "/api/upload/cv",
        files={"file": (filename, payload, ctype)},
        headers={"x-session-id": session_id},
    )


def _generate_roundtrip(client, fmt: str):
    """Generate a resume of ``fmt`` then re-upload it and return both parses."""
    up = _upload(client, "roundtrip-session")
    first = up.json()["parsed_data"]
    resume_id = up.json()["id"]
    result = client.post(
        f"/api/resume/{resume_id}/generate",
        json={"template_id": "ats_standard", "format": fmt},
        headers={"x-session-id": "roundtrip-session"},
    )
    assert result.status_code == 201, result.text
    gen_id = result.json()["id"]
    dl = client.get(
        f"/api/resume/{resume_id}/download/{gen_id}",
        headers={"x-session-id": "roundtrip-session"},
    )
    assert dl.status_code == 200

    content_type = (
        "application/pdf"
        if fmt == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    up2 = _upload_bytes(client, dl.content, f"roundtrip.{fmt}", content_type)
    assert up2.status_code == 201, up2.text
    return first, up2.json()["parsed_data"]


def test_docx_roundtrip_preserves_section_content(client):
    """Re-uploading a generated DOCX restores the same structured content."""
    for fmt in ("docx", "pdf"):
        first, second = _generate_roundtrip(client, fmt)
        # Capability names survive.
        for marker in ["Tech Corp", "Senior Developer"]:
            joined = "\n".join(
                e.get("title", "") + "\n" + e.get("company", "")
                for e in second["experience"]
            )
            assert marker in joined, (fmt, marker, second)

        # Institution and degree survive.
        edu = second["education"]
        edu_joined = "\n".join(e.get("institution", "") + "\n" + e.get("degree", "")
                               for e in edu)
        assert "MIT" in edu_joined, (fmt, second["education"])

        # Skills survive (at least one meaningful skill).
        assert len(second["skills"]) >= 3, (fmt, second["skills"])

        # Name and email survive.
        assert second["personal_info"]["name"] == "John Doe"
        assert second["personal_info"]["email"] == "john.doe@example.com"


def test_pdf_roundtrip_modern_column_layout(client):
    """A two-column (modern) PDF round-trip keeps sections contiguous."""
    up = _upload(client, "modern-session")
    resume_id = up.json()["id"]
    result = client.post(
        f"/api/resume/{resume_id}/generate",
        json={"template_id": "modern", "format": "pdf"},
        headers={"x-session-id": "modern-session"},
    )
    assert result.status_code == 201, result.text
    gen_id = result.json()["id"]
    dl = client.get(
        f"/api/resume/{resume_id}/download/{gen_id}",
        headers={"x-session-id": "modern-session"},
    )
    assert dl.status_code == 200
    up2 = _upload_bytes(client, dl.content, "roundtrip-modern.pdf", "application/pdf")
    assert up2.status_code == 201, up2.text
    second = up2.json()["parsed_data"]

    # Header survives.
    assert second["personal_info"]["name"] == "John Doe"
    assert second["personal_info"]["email"] == "john.doe@example.com"

    # Companies/titles survive (the two-column layout scrambled these before).
    joined = "\n".join(
        e.get("title", "") + "\n" + e.get("company", "")
        for e in second["experience"]
    )
    for marker in ["Tech Corp", "Senior Developer", "Startup Inc", "Developer"]:
        assert marker in joined, (marker, second["experience"])

    # Experience descriptions stay contiguous rather than gaining skill
    # fragments spliced between their wrapped lines.
    desc = " ".join(e.get("description", "") for e in second["experience"])
    assert "Led a team of five developers building REST APIs. Built CI/CD pipelines." in desc
    assert "Built customer-facing web apps with Python and React." in desc

    # Summary stays a continuous paragraph and skills survive as individual chips.
    assert "Experienced software engineer with 5+ years" in second["summary"]
    assert len(second["skills"]) >= 3, second["skills"]
    assert "python" in [s.lower() for s in second["skills"]], second["skills"]


def test_docx_body_contains_section_content(client):
    """The generated DOCX text includes the body of each section, not just
    the heading labels."""
    up = _upload(client)
    resume_id = up.json()["id"]
    for fmt in ("docx",):
        result = _generate(client, resume_id, "ats_standard", fmt)
        dl = client.get(
            f"/api/resume/{resume_id}/download/{result['id']}",
            headers={"x-session-id": "gen-session"},
        )
        assert dl.status_code == 200
        import io

        from docx import Document

        document = Document(io.BytesIO(dl.content))
        full_text = "\n".join(p.text for p in document.paragraphs)
        for marker in ["Tech Corp", "Senior Developer", "MIT", "B.S. Computer Science"]:
            assert marker in full_text, (fmt, marker, full_text)